from __future__ import annotations

import shutil
from pathlib import Path

import matplotlib

matplotlib.use("Agg")
import matplotlib.pyplot as plt
import pandas as pd
from docx import Document
from docx.shared import Inches


ROOT = Path(__file__).resolve().parents[2]
OUT = ROOT / "paper_results_discussion_pack_behavioral_focus"
TABLES = OUT / "tables"
FIGURES = OUT / "figures"
for d in (OUT, TABLES, FIGURES):
    d.mkdir(parents=True, exist_ok=True)


def rel(path: Path) -> str:
    return str(path.relative_to(ROOT)).replace("\\", "/")


def read_csv(path: Path) -> pd.DataFrame:
    return pd.read_csv(path)


def copy_if_exists(src: Path, dst: Path) -> bool:
    if src.exists():
        dst.parent.mkdir(parents=True, exist_ok=True)
        shutil.copy2(src, dst)
        return True
    return False


# Core evidence sources.
same_data_path = ROOT / "three_new_models_vs_old_delta3_non4060/results/tables/overall_ranking_with_feedbackaware_wsls.csv"
mixed_nonstick_path = ROOT / "nonstickiness_pattern_analysis/results/tables/nonstickiness_effect_summary.csv"
risky_summary_path = ROOT / "risky_choice_actual_main_effects_no_blocks.csv"
repetition_summary_path = ROOT / "repetition_actual_main_effects_no_blocks.csv"
integrated_path = ROOT / "initial_model_integrated_factors/results/tables/integrated_main_effects_contrasts.csv"
wsls_param_path = ROOT / "noisy_wsls_feedbackaware_non4060/results/tables/parameter_summary_noisy_wsls_feedbackaware_non4060.csv"
wsls_prob_path = ROOT / "noisy_wsls_feedbackaware_non4060/results/tables/feedback_condition_probabilities_noisy_wsls_feedbackaware_non4060.csv"
heldout_path = ROOT / "neural_synthetic_bayes_analysis/results/tables/model_comparison_train_test.csv"
heldout_cond_path = ROOT / "neural_synthetic_bayes_analysis/results/tables/nn_test_accuracy_by_condition.csv"
loo_path = ROOT / "neural_synthetic_bayes_analysis/results/tables/synthetic_bayes_train_loo_waic_overall.csv"

same = read_csv(same_data_path)
mixed = read_csv(mixed_nonstick_path)
risky_summary = read_csv(risky_summary_path)
repetition_summary = read_csv(repetition_summary_path)
integrated = read_csv(integrated_path)
wsls_prob = read_csv(wsls_prob_path)
heldout = read_csv(heldout_path)
heldout_cond = read_csv(heldout_cond_path)
loo = read_csv(loo_path)

for src in [
    same_data_path,
    mixed_nonstick_path,
    risky_summary_path,
    repetition_summary_path,
    integrated_path,
    wsls_param_path,
    wsls_prob_path,
    heldout_path,
    heldout_cond_path,
    loo_path,
]:
    copy_if_exists(src, TABLES / src.name)


def same_acc(model_id: str) -> float:
    return float(same.loc[same["model_id"] == model_id, "top1_weighted"].iloc[0])


def model_label(model_id: str) -> str:
    return str(same.loc[same["model_id"] == model_id, "model_label"].iloc[0])


def held_acc(model: str) -> float:
    return float(heldout.loc[heldout["model"] == model, "test_top1_accuracy"].iloc[0])


def risky_mean(variable: str, level: str) -> float:
    row = risky_summary[(risky_summary["Variable"] == variable) & (risky_summary["Level"] == level)].iloc[0]
    return float(row["Mean_Risky_Chance"])


def rep_mean(variable: str, level: str) -> float:
    row = repetition_summary[(repetition_summary["Variable"] == variable) & (repetition_summary["Level"] == level)].iloc[0]
    return float(row["Mean_Repetition"])


def int_row(parameter: str, effect: str) -> pd.Series:
    return integrated[(integrated["parameter"] == parameter) & (integrated["effect"] == effect)].iloc[0]


def mixed_row(analysis: str) -> pd.Series:
    return mixed[mixed["analysis"] == analysis].iloc[0]


def wsls_value(parameter: str, column: str = "mean") -> float:
    return float(wsls_prob.loc[wsls_prob["parameter"] == parameter, column].iloc[0])


focus_models = pd.DataFrame(
    [
        {
            "model_id": "old_original_benchmark",
            "model_family": "Original BMT / hierarchical softmax",
            "role": "theory-linked baseline",
            "top1_weighted": same_acc("old_original_benchmark"),
        },
        {
            "model_id": "modelB_decay",
            "model_family": "Model B: Decay-RL + Stickiness",
            "role": "main interpretable RL model",
            "top1_weighted": same_acc("modelB_decay"),
        },
        {
            "model_id": "noisy_wsls_feedbackaware",
            "model_family": "Feedback-aware Noisy WSLS",
            "role": "simple behavioral strategy model",
            "top1_weighted": same_acc("noisy_wsls_feedbackaware"),
        },
    ]
)
focus_models.to_csv(TABLES / "three_focus_model_comparison.csv", index=False)

fig, ax = plt.subplots(figsize=(8, 4.5))
plot_df = focus_models.sort_values("top1_weighted")
ax.barh(plot_df["model_family"], plot_df["top1_weighted"], color=["#8DA0CB", "#66C2A5", "#FC8D62"])
ax.axvline(0.5, color="black", linestyle="--", linewidth=1)
ax.set_xlim(0.45, 0.72)
ax.set_xlabel("Same-data weighted top-1 accuracy")
ax.set_title("Three focus models")
for i, value in enumerate(plot_df["top1_weighted"]):
    ax.text(value + 0.004, i, f"{value:.3f}", va="center")
fig.tight_layout()
fig.savefig(FIGURES / "fig1_three_focus_model_comparison.png", dpi=300)
plt.close(fig)

figure_sources = [
    (
        FIGURES / "fig1_three_focus_model_comparison.png",
        "fig1_three_focus_model_comparison",
        "Results: three-model comparison",
        "Same-data weighted top-1 accuracy for the three focus models: Original BMT, Decay-RL + Stickiness, and Feedback-aware WSLS.",
    ),
    (
        ROOT / "initial_model_integrated_factors/results/plots/integrated_main_effects_parameter_densities.png",
        "fig2_bmt_parameter_contrasts.png",
        "Results: BMT parameter contrasts",
        "Posterior parameter distributions for alpha, beta, and gamma in the BMT-family integrated hierarchical softmax model.",
    ),
    (
        ROOT / "noisy_wsls_feedbackaware_non4060/results/plots/feedbackaware_wsls_full_vs_partial_probabilities.png",
        "fig3_feedbackaware_wsls_probabilities.png",
        "Results: feedback-aware WSLS",
        "Feedback-aware win-stay and lose-shift probabilities under Full and Partial feedback.",
    ),
    (
        ROOT / "nonstickiness_pattern_analysis/results/plots/feedback_uncertainty_after_previous_choice.png",
        "fig4_feedback_uncertainty_after_previous_choice.png",
        "Results: mixed-effect behavioral validation",
        "Relative-uncertainty effect on risky choice after controlling previous choice.",
    ),
    (
        ROOT / "neural_synthetic_bayes_analysis/results/plots/model_comparison_train_test.png",
        "supp_fig1_heldout_model_comparison.png",
        "Supplementary: predictive validation",
        "Strict participant-level held-out accuracy for neural and Bayesian predictive validation models.",
    ),
    (
        ROOT / "neural_synthetic_bayes_analysis/results/plots/nn_test_accuracy_by_condition.png",
        "supp_fig2_nn_accuracy_by_condition.png",
        "Supplementary: predictive validation",
        "Held-out neural-network accuracy split by feedback, time pressure, and choice frame.",
    ),
]

figure_rows = []
for src, dest, section, caption in figure_sources:
    if src == FIGURES / "fig1_three_focus_model_comparison.png":
        pack_path = "figures/fig1_three_focus_model_comparison.png"
        source_path = rel(src)
        exists = src.exists()
    else:
        exists = copy_if_exists(src, FIGURES / dest)
        pack_path = f"figures/{dest}"
        source_path = rel(src)
    if exists:
        figure_rows.append(
            {
                "figure_id": dest.replace(".png", ""),
                "manuscript_section": section,
                "source_path": source_path,
                "pack_path": pack_path,
                "caption_draft": caption,
                "main_or_supplement": "supplement" if dest.startswith("supp_") else "main",
            }
        )
pd.DataFrame(figure_rows).to_csv(OUT / "figure_inventory_behavioral_focus.csv", index=False)


claims = []


def add_claim(claim_id, section, claim, statistic, value, source, row_filter, column, evidence, caveat=""):
    claims.append(
        {
            "claim_id": claim_id,
            "manuscript_section": section,
            "claim_text": claim,
            "statistic": statistic,
            "value": value,
            "source_file": rel(source),
            "source_row_or_filter": row_filter,
            "source_column": column,
            "evidence_level": evidence,
            "caveat": caveat,
        }
    )


add_claim("C01", "Behavioral summary", "Risky choice was descriptively higher under Full than Partial feedback.", "mean risky choice", f"Full={risky_mean('Feedback_Type','Full'):.3f}; Partial={risky_mean('Feedback_Type','Partial'):.3f}", risky_summary_path, "Variable == Feedback_Type", "Mean_Risky_Chance", "descriptive behavioral summary")
add_claim("C02", "Behavioral summary", "Risky choice was descriptively higher in Reject than Accept frames.", "mean risky choice", f"Accept={risky_mean('Choice_Frame','Accept'):.3f}; Reject={risky_mean('Choice_Frame','Reject'):.3f}", risky_summary_path, "Variable == Choice_Frame", "Mean_Risky_Chance", "descriptive behavioral summary")
add_claim("C03", "Behavioral summary", "Repetition was descriptively higher under time pressure than no time pressure.", "mean repetition", f"NoTP={rep_mean('Time_Pressure','No'):.3f}; TP={rep_mean('Time_Pressure','Yes'):.3f}", repetition_summary_path, "Variable == Time_Pressure", "Mean_Repetition", "descriptive behavioral summary")
add_claim("C04", "Behavioral summary", "Repetition was descriptively higher under Partial than Full feedback.", "mean repetition", f"Full={rep_mean('Feedback_Type','Full'):.3f}; Partial={rep_mean('Feedback_Type','Partial'):.3f}", repetition_summary_path, "Variable == Feedback_Type", "Mean_Repetition", "descriptive behavioral summary")

for analysis in [
    "Previous risky choice | stickiness control",
    "Full - Partial | relative uncertainty slope",
    "Accept - Reject | risky-choice bias",
    "Full feedback counterfactual regret predicts next switch",
    "Switch after regret/loss",
    "Partial feedback early-phase RU slope",
]:
    row = mixed_row(analysis)
    add_claim(
        f"C_mixed_{len(claims)+1:02d}",
        "Mixed-effect behavioral validation",
        analysis,
        "estimate, 95% CI, p-value",
        f"estimate={float(row['estimate']):.3f}; 95% CI [{float(row['conf_low']):.3f}, {float(row['conf_high']):.3f}]; p={float(row['p_value']):.3g}; significant={row['significant']}",
        mixed_nonstick_path,
        f"analysis == {analysis}",
        "estimate/conf_low/conf_high/p_value/significant",
        "mixed-effect / GLMM behavioral inference",
    )

add_claim("C_model_01", "Three-model comparison", "Model B improved same-data top-1 accuracy relative to the Original BMT benchmark.", "weighted top-1", f"Original BMT={same_acc('old_original_benchmark'):.3f}; Model B={same_acc('modelB_decay'):.3f}", same_data_path, "model_id in {old_original_benchmark, modelB_decay}", "top1_weighted", "same-data model comparison", "Same-data posterior-mean accuracy is optimistic relative to held-out validation.")
add_claim("C_model_02", "Three-model comparison", "Feedback-aware WSLS performed below Model B but above the original BMT benchmark.", "weighted top-1", f"WSLS={same_acc('noisy_wsls_feedbackaware'):.3f}; Model B={same_acc('modelB_decay'):.3f}; Original BMT={same_acc('old_original_benchmark'):.3f}", same_data_path, "model_id in {noisy_wsls_feedbackaware, modelB_decay, old_original_benchmark}", "top1_weighted", "same-data model comparison", "WSLS uses fewer mechanisms and is interpreted as a strategy model.")

for parameter, effect in [("gamma", "Time Pressure"), ("gamma", "Feedback"), ("alpha", "Feedback"), ("alpha", "Choice Frame"), ("beta", "Feedback")]:
    row = int_row(parameter, effect)
    add_claim(
        f"C_bmt_{parameter}_{effect.replace(' ', '_')}",
        "BMT parameter contrasts",
        f"The BMT-family integrated softmax model estimated a {effect.lower()} contrast for {parameter}; the 95% HDI {'excluded' if str(row['hdi95_crosses_zero']).upper() == 'FALSE' else 'crossed'} zero.",
        "posterior mean difference and 95% HDI",
        f"diff={float(row['diff_mean']):.3f}; 95% HDI [{float(row['diff_hdi95_lo']):.3f}, {float(row['diff_hdi95_hi']):.3f}]; crosses_zero={row['hdi95_crosses_zero']}",
        integrated_path,
        f"parameter == {parameter}; effect == {effect}",
        "diff_mean/diff_hdi95_lo/diff_hdi95_hi/hdi95_crosses_zero",
        "Bayesian posterior contrast",
        "Main-effects integrated model; task-specific heterogeneity is not fully represented.",
    )

add_claim("C_wsls_01", "Feedback-aware WSLS", "The feedback-aware WSLS model estimated lower win-stay probability under Full than Partial feedback.", "posterior probability difference", f"Full={wsls_value('p_win_stay_full'):.3f}; Partial={wsls_value('p_win_stay_partial'):.3f}; diff={wsls_value('diff_win_stay_full_minus_partial'):.3f}", wsls_prob_path, "p_win_stay_full, p_win_stay_partial, diff_win_stay_full_minus_partial", "mean/hdi95_low/hdi95_high", "Bayesian WSLS", "Feedback-aware win/loss coding differs by feedback condition.")
add_claim("C_wsls_02", "Feedback-aware WSLS", "The feedback-aware WSLS model estimated higher lose-shift probability under Full than Partial feedback.", "posterior probability difference", f"Full={wsls_value('p_lose_shift_full'):.3f}; Partial={wsls_value('p_lose_shift_partial'):.3f}; diff={wsls_value('diff_lose_shift_full_minus_partial'):.3f}", wsls_prob_path, "p_lose_shift_full, p_lose_shift_partial, diff_lose_shift_full_minus_partial", "mean/hdi95_low/hdi95_high", "Bayesian WSLS", "Feedback-aware win/loss coding differs by feedback condition.")
add_claim("C_supp_01", "Supplementary predictive validation", "The GRU achieved the highest strict held-out participant accuracy among supplementary predictive models.", "held-out top-1 accuracy", f"GRU={held_acc('NN direct GRU'):.3f}; MLP={held_acc('NN direct MLP'):.3f}; synthetic ModelB={held_acc('Bayes Model B synthetic_rep1'):.3f}", heldout_path, "NN direct GRU, NN direct MLP, Bayes Model B synthetic_rep1", "test_top1_accuracy", "supplementary held-out predictive validation", "Not central to main behavioral interpretation.")
pd.DataFrame(claims).to_csv(OUT / "result_claims_behavioral_focus.csv", index=False)

literature_rows = [
    ("Prospect theory and risk", "Frame/risk-bias discussion", "KahnemanTversky1979", "https://www.econometricsociety.org/publications/econometrica/browse/1979/03/01/prospect-theory-analysis-decision-under-risk"),
    ("Choice framing", "Accept/reject frame interpretation", "TverskyKahneman1981", "https://web.stanford.edu/~jlmcc/Presentations/tversky_kahneman_1981.pdf"),
    ("Prediction-error learning", "Learning-rate interpretation", "RescorlaWagner1972", "https://psycnet.apa.org/record/1973-01307-001"),
    ("Reinforcement learning", "Bandit and Q-learning model language", "SuttonBarto2018", "https://mitpress.ublish.com/book/reinforcement-learning-an-introduction-2"),
    ("Reinforcement learning in behavioral decisions", "Experience-based human choice modeling", "ErevRoth1998", "https://www.hbs.edu/faculty/Pages/item.aspx?num=3641"),
    ("Bayesian bandit modeling", "Human bandit model comparison", "Steyvers2009", "https://www.sciencedirect.com/science/article/pii/S0022249608001090"),
    ("Restless bandits and uncertainty", "Feedback/uncertainty discussion", "SpeekenbrinkKonstantinidis2015", "https://discovery.ucl.ac.uk/id/eprint/1470745/"),
    ("Behavioral exploration", "Directed/random exploration in humans", "Wilson2014", "https://pubmed.ncbi.nlm.nih.gov/25347535/"),
    ("Computational exploration algorithms", "Uncertainty bonus vs sampling interpretation", "Gershman2018", "https://pubmed.ncbi.nlm.nih.gov/29289795/"),
    ("Speed-accuracy tradeoff", "Behavioral interpretation of time-pressure effects", "Wickelgren1977", "https://www.sciencedirect.com/science/article/pii/0001691877900129"),
    ("Bayesian model comparison", "LOO/WAIC interpretation", "Vehtari2017", "https://arxiv.org/abs/1507.04544"),
]
pd.DataFrame(literature_rows, columns=["theme", "discussion_use", "citation_key", "source_url"]).to_csv(OUT / "literature_map_behavioral_focus.csv", index=False)

apa_refs = """Erev, I., & Roth, A. E. (1998). Predicting how people play games: Reinforcement learning in experimental games with unique, mixed strategy equilibria. *American Economic Review, 88*(4), 848-881.

Gershman, S. J. (2018). Deconstructing the human algorithms for exploration. *Cognition, 173*, 34-42. https://doi.org/10.1016/j.cognition.2017.12.014

Kahneman, D., & Tversky, A. (1979). Prospect theory: An analysis of decision under risk. *Econometrica, 47*(2), 263-291. https://doi.org/10.2307/1914185

Rescorla, R. A., & Wagner, A. R. (1972). A theory of Pavlovian conditioning: Variations in the effectiveness of reinforcement and nonreinforcement. In A. H. Black & W. F. Prokasy (Eds.), *Classical conditioning II: Current research and theory* (pp. 64-99). Appleton-Century-Crofts.

Speekenbrink, M., & Konstantinidis, E. (2015). Uncertainty and exploration in a restless bandit problem. *Topics in Cognitive Science, 7*(2), 351-367. https://doi.org/10.1111/tops.12145

Steyvers, M., Lee, M. D., & Wagenmakers, E.-J. (2009). A Bayesian analysis of human decision-making on bandit problems. *Journal of Mathematical Psychology, 53*(3), 168-179. https://doi.org/10.1016/j.jmp.2008.11.002

Sutton, R. S., & Barto, A. G. (2018). *Reinforcement learning: An introduction* (2nd ed.). MIT Press.

Tversky, A., & Kahneman, D. (1981). The framing of decisions and the psychology of choice. *Science, 211*(4481), 453-458. https://doi.org/10.1126/science.7455683

Vehtari, A., Gelman, A., & Gabry, J. (2017). Practical Bayesian model evaluation using leave-one-out cross-validation and WAIC. *Statistics and Computing, 27*, 1413-1432. https://doi.org/10.1007/s11222-016-9696-4

Wickelgren, W. A. (1977). Speed-accuracy tradeoff and information processing dynamics. *Acta Psychologica, 41*(1), 67-85. https://doi.org/10.1016/0001-6918(77)90012-9

Wilson, R. C., Geana, A., White, J. M., Ludvig, E. A., & Cohen, J. D. (2014). Humans use directed and random exploration to solve the explore-exploit dilemma. *Journal of Experimental Psychology: General, 143*(6), 2074-2081. https://doi.org/10.1037/a0038199
"""
(OUT / "references_apa_behavioral_focus.txt").write_text(apa_refs, encoding="utf-8")

bib = """@article{ErevRoth1998, author={Erev, Ido and Roth, Alvin E.}, title={Predicting How People Play Games: Reinforcement Learning in Experimental Games with Unique, Mixed Strategy Equilibria}, journal={American Economic Review}, year={1998}, volume={88}, number={4}, pages={848--881}}
@article{Gershman2018, author={Gershman, Samuel J.}, title={Deconstructing the Human Algorithms for Exploration}, journal={Cognition}, year={2018}, volume={173}, pages={34--42}, doi={10.1016/j.cognition.2017.12.014}}
@article{KahnemanTversky1979, author={Kahneman, Daniel and Tversky, Amos}, title={Prospect Theory: An Analysis of Decision under Risk}, journal={Econometrica}, year={1979}, volume={47}, number={2}, pages={263--291}, doi={10.2307/1914185}}
@incollection{RescorlaWagner1972, author={Rescorla, Robert A. and Wagner, Allan R.}, title={A Theory of Pavlovian Conditioning: Variations in the Effectiveness of Reinforcement and Nonreinforcement}, booktitle={Classical Conditioning II: Current Research and Theory}, editor={Black, Abraham H. and Prokasy, William F.}, publisher={Appleton-Century-Crofts}, address={New York}, year={1972}, pages={64--99}}
@article{SpeekenbrinkKonstantinidis2015, author={Speekenbrink, Maarten and Konstantinidis, Emmanouil}, title={Uncertainty and Exploration in a Restless Bandit Problem}, journal={Topics in Cognitive Science}, year={2015}, volume={7}, number={2}, pages={351--367}, doi={10.1111/tops.12145}}
@article{Steyvers2009, author={Steyvers, Mark and Lee, Michael D. and Wagenmakers, Eric-Jan}, title={A Bayesian Analysis of Human Decision-Making on Bandit Problems}, journal={Journal of Mathematical Psychology}, year={2009}, volume={53}, number={3}, pages={168--179}, doi={10.1016/j.jmp.2008.11.002}}
@book{SuttonBarto2018, author={Sutton, Richard S. and Barto, Andrew G.}, title={Reinforcement Learning: An Introduction}, edition={2}, publisher={MIT Press}, year={2018}}
@article{TverskyKahneman1981, author={Tversky, Amos and Kahneman, Daniel}, title={The Framing of Decisions and the Psychology of Choice}, journal={Science}, year={1981}, volume={211}, number={4481}, pages={453--458}, doi={10.1126/science.7455683}}
@article{Vehtari2017, author={Vehtari, Aki and Gelman, Andrew and Gabry, Jonah}, title={Practical Bayesian Model Evaluation Using Leave-One-Out Cross-Validation and WAIC}, journal={Statistics and Computing}, year={2017}, volume={27}, pages={1413--1432}, doi={10.1007/s11222-016-9696-4}}
@article{Wickelgren1977, author={Wickelgren, Wayne A.}, title={Speed-Accuracy Tradeoff and Information Processing Dynamics}, journal={Acta Psychologica}, year={1977}, volume={41}, number={1}, pages={67--85}, doi={10.1016/0001-6918(77)90012-9}}
@article{Wilson2014, author={Wilson, Robert C. and Geana, Andra and White, John M. and Ludvig, Elliot A. and Cohen, Jonathan D.}, title={Humans Use Directed and Random Exploration to Solve the Explore--Exploit Dilemma}, journal={Journal of Experimental Psychology: General}, year={2014}, volume={143}, number={6}, pages={2074--2081}, doi={10.1037/a0038199}}
"""
(OUT / "references_behavioral_focus.bib").write_text(bib, encoding="utf-8")

results_text = f"""# Results

## Behavioral Overview

The analyses focused on eight non-4060 follow-up bandit tasks. The behavioral summaries showed that risky choice was higher under Full feedback than Partial feedback (Full = {risky_mean('Feedback_Type','Full'):.3f}; Partial = {risky_mean('Feedback_Type','Partial'):.3f}) and higher in Reject than Accept frames (Reject = {risky_mean('Choice_Frame','Reject'):.3f}; Accept = {risky_mean('Choice_Frame','Accept'):.3f}). Repetition also varied by condition: choices were repeated more often under time pressure than no time pressure (TP = {rep_mean('Time_Pressure','Yes'):.3f}; NoTP = {rep_mean('Time_Pressure','No'):.3f}) and more often under Partial than Full feedback (Partial = {rep_mean('Feedback_Type','Partial'):.3f}; Full = {rep_mean('Feedback_Type','Full'):.3f}). These descriptive results motivated the computational analyses by showing that feedback, frame, and time pressure were behaviorally relevant, but they did not by themselves identify the underlying mechanism.

## Mixed-Effect Behavioral Validation

The mixed-effect analyses provided direct behavioral evidence that choice history was important. After controlling for task structure, previous risky choice significantly predicted current risky choice (estimate = {float(mixed_row('Previous risky choice | stickiness control')['estimate']):.3f}, 95% CI [{float(mixed_row('Previous risky choice | stickiness control')['conf_low']):.3f}, {float(mixed_row('Previous risky choice | stickiness control')['conf_high']):.3f}], p = {float(mixed_row('Previous risky choice | stickiness control')['p_value']):.3g}). This supports the interpretation that repetition or stickiness was not merely a computational-model artifact; it was visible in the behavioral data.

However, reliable non-stickiness effects also remained. The Full - Partial contrast in the relative-uncertainty slope was significant after accounting for previous choice (estimate = {float(mixed_row('Full - Partial | relative uncertainty slope')['estimate']):.3f}, 95% CI [{float(mixed_row('Full - Partial | relative uncertainty slope')['conf_low']):.3f}, {float(mixed_row('Full - Partial | relative uncertainty slope')['conf_high']):.3f}], p = {float(mixed_row('Full - Partial | relative uncertainty slope')['p_value']):.3g}). Choice frame also predicted risky-choice bias (Accept - Reject estimate = {float(mixed_row('Accept - Reject | risky-choice bias')['estimate']):.3f}, 95% CI [{float(mixed_row('Accept - Reject | risky-choice bias')['conf_low']):.3f}, {float(mixed_row('Accept - Reject | risky-choice bias')['conf_high']):.3f}], p = {float(mixed_row('Accept - Reject | risky-choice bias')['p_value']):.3g}). These results indicate that feedback and frame effects were not reducible to simple repetition, although their strongest expression was behavioral rather than a stable learning-rate shift.

Switching behavior further supported a feedback-dependent strategy account. Regret or loss after the previous trial increased switching (estimate = {float(mixed_row('Switch after regret/loss')['estimate']):.3f}, 95% CI [{float(mixed_row('Switch after regret/loss')['conf_low']):.3f}, {float(mixed_row('Switch after regret/loss')['conf_high']):.3f}], p = {float(mixed_row('Switch after regret/loss')['p_value']):.3g}). Within Full feedback, counterfactual regret predicted next-trial switching (estimate = {float(mixed_row('Full feedback counterfactual regret predicts next switch')['estimate']):.3f}, 95% CI [{float(mixed_row('Full feedback counterfactual regret predicts next switch')['conf_low']):.3f}, {float(mixed_row('Full feedback counterfactual regret predicts next switch')['conf_high']):.3f}], p = {float(mixed_row('Full feedback counterfactual regret predicts next switch')['p_value']):.3g}). In Partial feedback, early relative uncertainty predicted risky choice (estimate = {float(mixed_row('Partial feedback early-phase RU slope')['estimate']):.3f}, 95% CI [{float(mixed_row('Partial feedback early-phase RU slope')['conf_low']):.3f}, {float(mixed_row('Partial feedback early-phase RU slope')['conf_high']):.3f}], p = {float(mixed_row('Partial feedback early-phase RU slope')['p_value']):.3g}). Thus, after controlling for stickiness, the strongest remaining behavioral signals involved uncertainty and feedback-dependent outcome evaluation.

## Three Focus Models

The main model comparison focused on three interpretable models. The Original BMT / hierarchical softmax model served as the theory-linked baseline, Model B represented a dynamic Decay-RL model with stickiness, and the feedback-aware WSLS model represented a simple behavioral strategy account. In same-data weighted top-1 accuracy, Model B achieved {same_acc('modelB_decay'):.3f}, the feedback-aware WSLS model achieved {same_acc('noisy_wsls_feedbackaware'):.3f}, and the Original BMT benchmark achieved {same_acc('old_original_benchmark'):.3f}. Model B therefore improved prediction relative to the original benchmark, while WSLS remained useful as a simple behavioral explanation despite lower accuracy than Model B.

The BMT-family integrated parameter analysis showed that gamma was the clearest factor-sensitive parameter. The Time Pressure contrast for gamma was credibly positive (TP - NoTP = {float(int_row('gamma','Time Pressure')['diff_mean']):.3f}, 95% HDI [{float(int_row('gamma','Time Pressure')['diff_hdi95_lo']):.3f}, {float(int_row('gamma','Time Pressure')['diff_hdi95_hi']):.3f}], HDI excluded zero), consistent with stronger perseveration under time pressure. The Feedback contrast for gamma also excluded zero (Full - Partial = {float(int_row('gamma','Feedback')['diff_mean']):.3f}, 95% HDI [{float(int_row('gamma','Feedback')['diff_hdi95_lo']):.3f}, {float(int_row('gamma','Feedback')['diff_hdi95_hi']):.3f}]), consistent with stronger stickiness under Partial than Full feedback in this coding. By comparison, alpha and beta showed feedback-related contrasts, but these learning/value-sensitivity effects were less consistent across model variants and should be interpreted more cautiously.

The feedback-aware WSLS model clarified the direction of the feedback effect in strategic terms. Win-stay probability was lower under Full than Partial feedback (Full = {wsls_value('p_win_stay_full'):.3f}; Partial = {wsls_value('p_win_stay_partial'):.3f}; Full - Partial = {wsls_value('diff_win_stay_full_minus_partial'):.3f}, 95% HDI [{wsls_value('diff_win_stay_full_minus_partial','hdi95_low'):.3f}, {wsls_value('diff_win_stay_full_minus_partial','hdi95_high'):.3f}]). Lose-shift probability was higher under Full than Partial feedback (Full = {wsls_value('p_lose_shift_full'):.3f}; Partial = {wsls_value('p_lose_shift_partial'):.3f}; Full - Partial = {wsls_value('diff_lose_shift_full_minus_partial'):.3f}, 95% HDI [{wsls_value('diff_lose_shift_full_minus_partial','hdi95_low'):.3f}, {wsls_value('diff_lose_shift_full_minus_partial','hdi95_high'):.3f}]). This pattern suggests that Partial feedback promoted repetition following favorable chosen outcomes, whereas Full feedback promoted switching when counterfactual information indicated that the unchosen option was better.

## Supplementary Predictive Validation

Neural-network and synthetic-Bayesian analyses were treated as supplementary predictive checks rather than central behavioral models. Under strict participant-level held-out validation, the GRU achieved 0.654 top-1 accuracy, the MLP achieved 0.626, and Model B trained on synthetic neural-network choices achieved 0.600. These results suggest that sequence models captured additional predictive history structure, but synthetic choices did not replace real participant data for fitting an interpretable Bayesian RL model.
"""

discussion_text = """# Discussion

The behavioral and computational results converge on a simple conclusion: the most robust structure in the data was choice-history dependence combined with feedback-dependent strategy adjustment. The results do not support a single-mechanism story in which feedback, time pressure, and choice frame all map cleanly onto one learning-rate parameter. Instead, participants' behavior was best understood as a combination of reinforcement learning, repetition, uncertainty sensitivity, and response bias.

The clearest computational signal was stickiness. Both the mixed-effect behavioral analyses and the BMT-family parameter contrasts indicated that recent choice history strongly shaped current choice. This is consistent with reinforcement-learning approaches in which choices are not determined only by expected value, but also by action tendencies, perseveration, or response propensities learned through experience (Erev & Roth, 1998; Sutton & Barto, 2018). In the present data, time pressure increased the gamma/stickiness parameter and descriptively increased repetition. This pattern suggests that time pressure made participants rely more on recent actions, consistent with behavioral speed-accuracy tradeoff accounts in which time constraints reduce the opportunity for slower comparison processes (Wickelgren, 1977).

Feedback changed behavior primarily by changing the information available for learning and switching. Full feedback allowed participants to compare chosen and unchosen outcomes, whereas Partial feedback only provided chosen-outcome information. This distinction is central in bandit tasks because feedback determines what can be learned from unchosen options and how uncertainty is resolved (Steyvers et al., 2009; Speekenbrink & Konstantinidis, 2015). The mixed-effect analyses showed that relative uncertainty and counterfactual regret remained predictive after controlling for previous choice, and the WSLS model showed higher lose-shift under Full feedback but higher win-stay under Partial feedback. Thus, Full feedback appears to support counterfactual switching, whereas Partial feedback supports stronger repetition after favorable chosen outcomes.

Choice frame was better explained as a risk-bias or response-bias effect than as a stable shift in learning. Classic framing work shows that equivalent choice problems can produce different risk preferences depending on how outcomes are described (Kahneman & Tversky, 1979; Tversky & Kahneman, 1981). In this dataset, frame effects appeared clearly in risky-choice bias in the mixed-effect analyses, but they were not consistently recovered as robust beta or gamma shifts in the computational models. This suggests that Accept/Reject framing may alter the mapping between task representation and risky response rather than fundamentally changing how participants update values.

The uncertainty results support a behavioral bandit interpretation rather than a purely accuracy-based model-comparison interpretation. Human bandit behavior often reflects a mixture of directed exploration, random exploration, and heuristic strategies rather than a single optimal policy (Gershman, 2018; Wilson et al., 2014). Here, uncertainty effects were most visible after controlling for previous choice and separating feedback conditions. This means that uncertainty-based behavior was present, but it was partly masked by strong repetition. Future models should therefore include both uncertainty-sensitive choice terms and explicit choice-history terms, rather than treating stickiness as a nuisance parameter only.

Model comparison was useful, but the highest-accuracy model should not automatically be treated as the best psychological explanation. Model B improved prediction over the Original BMT benchmark and retained interpretable parameters for learning, value sensitivity, decay, and stickiness. The WSLS model was less accurate but provided a transparent account of how feedback changed simple strategy use. This illustrates why model comparison should balance predictive performance with interpretability, especially in cognitive bandit tasks. Predictive criteria such as LOO/WAIC are valuable for model evaluation, but they do not by themselves identify the true psychological mechanism (Vehtari et al., 2017).

The main limitation is that several parameter-level effects remain model-dependent. Feedback effects on alpha and beta appeared in the BMT-family model, but their interpretation was less stable than the gamma/stickiness effects. In addition, some task-level parameter comparisons were descriptive rather than full posterior contrasts, so they should not be overinterpreted. A second limitation is that time pressure was analyzed behaviorally and computationally, but the present task was not optimized to isolate speed-accuracy mechanisms. Future follow-up studies should strengthen the time-pressure manipulation, include explicit manipulation checks, and estimate participant-level random slopes for repetition, uncertainty sensitivity, and frame bias.

Overall, the results suggest that the follow-up bandit task successfully captured meaningful behavioral structure. The strongest pattern was not a general increase or decrease in learning, but a shift in how participants combined recent choice history, feedback information, uncertainty, and frame-dependent response tendencies. A concise interpretation is that feedback shaped what participants could learn from outcomes, time pressure increased reliance on repeated actions, and choice frame shifted risky-choice bias.
"""

full_md = results_text + "\n\n" + discussion_text + "\n\n# References\n\n" + apa_refs
(OUT / "results_discussion_behavioral_focus.md").write_text(full_md, encoding="utf-8")


def add_markdown_to_doc(doc: Document, text: str) -> None:
    for block in text.strip().split("\n\n"):
        if not block.strip():
            continue
        block = block.strip()
        if block.startswith("# "):
            doc.add_heading(block[2:], level=1)
        elif block.startswith("## "):
            doc.add_heading(block[3:], level=2)
        else:
            doc.add_paragraph(block.replace("\n", " "))


doc = Document()
doc.add_heading("Behavioral-Focused Results and Discussion Draft", level=0)
doc.add_paragraph("Main text focuses on Original BMT, Model B, Feedback-aware WSLS, and mixed-effect behavioral validation. Neural/synthetic analyses are retained only as supplementary predictive checks.")
add_markdown_to_doc(doc, results_text)
doc.add_heading("Selected Figures", level=1)
for row in figure_rows:
    if row["main_or_supplement"] != "main":
        continue
    doc.add_heading(row["figure_id"], level=2)
    fig_path = OUT / row["pack_path"]
    if fig_path.exists():
        doc.add_picture(str(fig_path), width=Inches(6.2))
    doc.add_paragraph(row["caption_draft"])
add_markdown_to_doc(doc, discussion_text)
doc.add_heading("Supplementary Predictive Validation Figures", level=1)
for row in figure_rows:
    if row["main_or_supplement"] != "supplement":
        continue
    doc.add_heading(row["figure_id"], level=2)
    fig_path = OUT / row["pack_path"]
    if fig_path.exists():
        doc.add_picture(str(fig_path), width=Inches(6.2))
    doc.add_paragraph(row["caption_draft"])
doc.add_heading("References", level=1)
for ref in apa_refs.strip().split("\n\n"):
    doc.add_paragraph(ref)
doc.save(OUT / "results_discussion_behavioral_focus.docx")

readme = f"""# Behavioral-Focused Paper Pack

This pack rewrites the Results and Discussion around behavioral bandit modeling.

## Read first

1. `results_discussion_behavioral_focus.md` or `.docx`
2. `result_claims_behavioral_focus.csv`
3. `figure_inventory_behavioral_focus.csv`
4. `references_apa_behavioral_focus.txt`

## Main-text model focus

- Original BMT / hierarchical softmax benchmark
- Model B: Decay-RL + Stickiness
- Feedback-aware Noisy WSLS

## Supplementary only

Neural-network and synthetic-Bayesian results are included only as predictive validation, not as central behavioral evidence.

## Citation scope

Main text uses behavioral decision-making, bandit-task, and computational modeling references only. Neuroscience-focused citations were intentionally removed.

## Evidence rule

Every numeric claim in the draft is indexed in `result_claims_behavioral_focus.csv`.
"""
(OUT / "chatgpt_readme_behavioral_focus.md").write_text(readme, encoding="utf-8")

print(f"Created behavioral-focused paper pack: {OUT}")
